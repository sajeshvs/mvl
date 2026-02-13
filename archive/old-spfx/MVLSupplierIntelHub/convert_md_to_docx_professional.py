"""
Professional Markdown to DOCX Converter
Converts DETAILED_IMPROVEMENT_NOTES.md to a beautifully formatted Word document
with MVL branding and professional design.

Author: MVL Data Engineering Team
Date: February 10, 2026
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os


class ProfessionalDocumentDesigner:
    """
    Designs a professional document with MVL branding.
    
    Design Philosophy:
    - Clean, modern layout with ample white space
    - Consistent color scheme (MVL brand colors)
    - Hierarchical typography for easy scanning
    - Professional tables with alternating row colors
    - Visual separators between sections
    - Emphasis on readability and professionalism
    """
    
    # MVL Brand Colors (Professional Blue/Gray scheme)
    COLOR_PRIMARY = RGBColor(0, 51, 102)      # Deep Blue - Main headings
    COLOR_SECONDARY = RGBColor(0, 102, 204)   # Bright Blue - Sub headings
    COLOR_ACCENT = RGBColor(51, 153, 204)     # Light Blue - Accents
    COLOR_TEXT = RGBColor(51, 51, 51)         # Dark Gray - Body text
    COLOR_TABLE_HEADER = RGBColor(0, 51, 102) # Blue - Table headers
    COLOR_TABLE_ALT = RGBColor(242, 242, 242) # Light Gray - Alternate rows
    COLOR_SUCCESS = RGBColor(34, 139, 34)     # Green - Success indicators
    COLOR_CODE_BG = RGBColor(245, 245, 245)   # Light Gray - Code blocks
    
    def __init__(self, logo_path):
        """
        Initialize the document with professional settings.
        
        Args:
            logo_path: Path to MVL logo file
        """
        self.doc = Document()
        self.logo_path = logo_path
        
        # Set document margins (narrower for more content)
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
            section.header_distance = Cm(1.0)
            section.footer_distance = Cm(1.0)
        
        # Configure professional styles
        self._setup_styles()
        
        # Add header with logo
        self._create_header()
        
        # Add footer with page numbers
        self._create_footer()
    
    def _setup_styles(self):
        """
        Create professional, consistent styles for the entire document.
        
        Style Hierarchy:
        - Title: Large, bold, primary color - for main document title
        - Heading 1: Major sections, centered, primary color
        - Heading 2: Subsections, centered, secondary color
        - Heading 3: Minor sections, left-aligned, accent color
        - Normal: Body text, comfortable reading size
        - Code: Monospace, gray background
        - Bullet: Formatted bullet points
        """
        styles = self.doc.styles
        
        # Title Style (Document Title)
        title_style = styles['Title']
        title_font = title_style.font
        title_font.name = 'Calibri Light'
        title_font.size = Pt(32)
        title_font.bold = True
        title_font.color.rgb = self.COLOR_PRIMARY
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(20)
        title_style.paragraph_format.space_before = Pt(10)
        
        # Heading 1 Style (Main Sections)
        h1_style = styles['Heading 1']
        h1_font = h1_style.font
        h1_font.name = 'Calibri'
        h1_font.size = Pt(20)
        h1_font.bold = True
        h1_font.color.rgb = self.COLOR_PRIMARY
        h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h1_style.paragraph_format.space_before = Pt(20)
        h1_style.paragraph_format.space_after = Pt(12)
        h1_style.paragraph_format.keep_with_next = True
        
        # Heading 2 Style (Subsections)
        h2_style = styles['Heading 2']
        h2_font = h2_style.font
        h2_font.name = 'Calibri'
        h2_font.size = Pt(16)
        h2_font.bold = True
        h2_font.color.rgb = self.COLOR_SECONDARY
        h2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h2_style.paragraph_format.space_before = Pt(16)
        h2_style.paragraph_format.space_after = Pt(10)
        h2_style.paragraph_format.keep_with_next = True
        
        # Heading 3 Style (Minor sections)
        h3_style = styles['Heading 3']
        h3_font = h3_style.font
        h3_font.name = 'Calibri'
        h3_font.size = Pt(14)
        h3_font.bold = True
        h3_font.color.rgb = self.COLOR_ACCENT
        h3_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h3_style.paragraph_format.space_before = Pt(12)
        h3_style.paragraph_format.space_after = Pt(8)
        h3_style.paragraph_format.left_indent = Pt(0)
        
        # Normal Style (Body Text)
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Calibri'
        normal_font.size = Pt(11)
        normal_font.color.rgb = self.COLOR_TEXT
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # List Bullet Style
        try:
            list_style = styles['List Bullet']
            list_font = list_style.font
            list_font.name = 'Calibri'
            list_font.size = Pt(11)
            list_font.color.rgb = self.COLOR_TEXT
            list_style.paragraph_format.space_after = Pt(3)
            list_style.paragraph_format.left_indent = Inches(0.5)
        except KeyError:
            pass
    
    def _create_header(self):
        """
        Create professional header with centered MVL logo.
        
        Design:
        - Centered logo (scaled appropriately)
        - Bottom border for separation
        - Consistent across all pages
        """
        section = self.doc.sections[0]
        header = section.header
        
        # Clear existing header content
        header.is_linked_to_previous = False
        
        # Add centered logo
        if os.path.exists(self.logo_path):
            header_para = header.paragraphs[0]
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add logo with appropriate sizing
            run = header_para.add_run()
            run.add_picture(self.logo_path, width=Inches(2.0))
            
            # Add spacing after logo
            header_para.paragraph_format.space_after = Pt(8)
            
            # Add bottom border to header
            self._add_paragraph_border(header_para, bottom=True)
    
    def _create_footer(self):
        """
        Create professional footer with page numbers and document info.
        
        Design:
        - Page numbers on right
        - Document title on left
        - Top border for separation
        """
        section = self.doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        
        # Create footer paragraph
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add top border
        self._add_paragraph_border(footer_para, top=True)
        
        # Add page number
        run = footer_para.add_run()
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Add page number field
        footer_para.add_run("Page ")
        self._add_page_number(footer_para)
        footer_para.add_run(" | MVL Supply Chain Intel Hub - Data Improvement Report")
    
    def _add_page_number(self, paragraph):
        """Add dynamic page number field to paragraph."""
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
    
    def _add_paragraph_border(self, paragraph, top=False, bottom=False):
        """Add border to paragraph for visual separation."""
        p = paragraph._element
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        
        if top:
            top_border = OxmlElement('w:top')
            top_border.set(qn('w:val'), 'single')
            top_border.set(qn('w:sz'), '6')
            top_border.set(qn('w:space'), '1')
            top_border.set(qn('w:color'), '336699')
            pBdr.append(top_border)
        
        if bottom:
            bottom_border = OxmlElement('w:bottom')
            bottom_border.set(qn('w:val'), 'single')
            bottom_border.set(qn('w:sz'), '6')
            bottom_border.set(qn('w:space'), '1')
            bottom_border.set(qn('w:color'), '336699')
            pBdr.append(bottom_border)
        
        pPr.append(pBdr)
    
    def add_title_page(self, title, subtitle, date, version):
        """
        Create an attractive title page.
        
        Args:
            title: Main document title
            subtitle: Document subtitle/description
            date: Document date
            version: Document version
        """
        # Add vertical space
        for _ in range(5):
            self.doc.add_paragraph()
        
        # Main title
        title_para = self.doc.add_paragraph(title, style='Title')
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle_para = self.doc.add_paragraph(subtitle)
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.runs[0]
        subtitle_run.font.size = Pt(16)
        subtitle_run.font.color.rgb = self.COLOR_SECONDARY
        subtitle_run.font.italic = True
        
        # Add spacing
        self.doc.add_paragraph()
        
        # Create info table (centered)
        info_para = self.doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        info_text = f"Version: {version}\nDate: {date}\nTotal Records Processed: 17,864"
        info_para.add_run(info_text).font.size = Pt(12)
        
        # Add spacing
        for _ in range(3):
            self.doc.add_paragraph()
        
        # Add decorative line
        line_para = self.doc.add_paragraph("─" * 80)
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run = line_para.runs[0]
        line_run.font.color.rgb = self.COLOR_ACCENT
        
        # Page break for content
        self.doc.add_page_break()
    
    def add_heading(self, text, level=1):
        """
        Add a professionally formatted heading.
        
        Args:
            text: Heading text
            level: Heading level (1, 2, or 3)
        """
        # Clean heading text (remove markdown markers)
        clean_text = re.sub(r'^#+\s*', '', text).strip()
        clean_text = re.sub(r'^\d+\.\s*', '', clean_text)
        
        # Add heading
        style_name = f'Heading {level}'
        heading = self.doc.add_paragraph(clean_text, style=style_name)
        
        # Add visual separator after major headings
        if level == 1:
            separator = self.doc.add_paragraph("━" * 100)
            separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sep_run = separator.runs[0]
            sep_run.font.size = Pt(8)
            sep_run.font.color.rgb = self.COLOR_ACCENT
            self.doc.add_paragraph()  # Add spacing
        
        return heading
    
    def add_paragraph(self, text, bold=False, italic=False, color=None):
        """
        Add a formatted paragraph.
        
        Args:
            text: Paragraph text
            bold: Make text bold
            italic: Make text italic
            color: Text color (RGBColor)
        """
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        
        if bold:
            run.font.bold = True
        if italic:
            run.font.italic = True
        if color:
            run.font.color.rgb = color
        
        return para
    
    def add_bullet_point(self, text, checked=False):
        """
        Add a formatted bullet point with optional checkmark.
        
        Args:
            text: Bullet text
            checked: Add checkmark indicator
        """
        para = self.doc.add_paragraph(style='List Bullet')
        
        if checked:
            # Add checkmark symbol
            run = para.add_run("✅ ")
            run.font.size = Pt(11)
        
        run = para.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = self.COLOR_TEXT
        
        return para
    
    def add_code_block(self, code):
        """
        Add a formatted code block with gray background.
        
        Args:
            code: Code text (can be multi-line)
        """
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        
        run = para.add_run(code)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(80, 80, 80)
        
        # Add background color (through shading)
        self._add_shading(para, self.COLOR_CODE_BG)
        
        return para
    
    def _add_shading(self, paragraph, color):
        """Add background shading to paragraph.
        
        Args:
            color: RGBColor object or hex string
        """
        shading_elm = OxmlElement('w:shd')
        # RGBColor converts to hex string directly via str()
        hex_color = str(color) if not isinstance(color, str) else color
        shading_elm.set(qn('w:fill'), hex_color)
        paragraph._element.get_or_add_pPr().append(shading_elm)
    
    def add_table(self, headers, rows):
        """
        Add a professionally formatted table.
        
        Design:
        - Blue header with white text
        - Alternating row colors (white/light gray)
        - Proper borders
        - Centered text
        
        Args:
            headers: List of header strings
            rows: List of row data (list of lists)
        """
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Format header row
        header_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            cell = header_cells[idx]
            cell.text = header
            
            # Header formatting
            cell_para = cell.paragraphs[0]
            cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_run = cell_para.runs[0]
            cell_run.font.bold = True
            cell_run.font.size = Pt(11)
            cell_run.font.color.rgb = RGBColor(255, 255, 255)
            
            # Header background color
            self._set_cell_background(cell, self.COLOR_TABLE_HEADER)
        
        # Add data rows
        for row_idx, row_data in enumerate(rows):
            row_cells = table.add_row().cells
            
            for idx, cell_data in enumerate(row_data):
                cell = row_cells[idx]
                cell.text = str(cell_data)
                
                # Cell formatting
                cell_para = cell.paragraphs[0]
                cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell_run = cell_para.runs[0]
                cell_run.font.size = Pt(10)
                
                # Alternate row colors
                if row_idx % 2 == 1:
                    self._set_cell_background(cell, self.COLOR_TABLE_ALT)
        
        # Add spacing after table
        self.doc.add_paragraph()
        
        return table
    
    def _set_cell_background(self, cell, color):
        """Set cell background color.
        
        Args:
            color: RGBColor object or hex string
        """
        shading_elm = OxmlElement('w:shd')
        # RGBColor converts to hex string directly via str()
        hex_color = str(color) if not isinstance(color, str) else color
        shading_elm.set(qn('w:fill'), hex_color)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    def add_info_box(self, title, content, icon="📊"):
        """
        Add an information box with colored border.
        
        Args:
            title: Box title
            content: Box content
            icon: Icon to display
        """
        # Add spacing
        self.doc.add_paragraph()
        
        # Title with icon
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run(f"{icon} {title}")
        title_run.font.bold = True
        title_run.font.size = Pt(12)
        title_run.font.color.rgb = self.COLOR_SECONDARY
        
        # Content
        content_para = self.doc.add_paragraph(content)
        content_para.paragraph_format.left_indent = Inches(0.3)
        content_run = content_para.runs[0]
        content_run.font.size = Pt(11)
        
        # Add spacing
        self.doc.add_paragraph()
    
    def add_before_after_comparison(self, before_title, before_data, after_title, after_data):
        """
        Add a before/after comparison section.
        
        Args:
            before_title: Title for before section
            before_data: Before data (dict or string)
            after_title: Title for after section
            after_data: After data (dict or string)
        """
        # Before section
        before_para = self.doc.add_paragraph()
        before_run = before_para.add_run(before_title)
        before_run.font.bold = True
        before_run.font.color.rgb = RGBColor(204, 0, 0)  # Red for "before"
        
        before_code = self.add_code_block(str(before_data))
        
        # After section
        after_para = self.doc.add_paragraph()
        after_run = after_para.add_run(after_title)
        after_run.font.bold = True
        after_run.font.color.rgb = self.COLOR_SUCCESS  # Green for "after"
        
        after_code = self.add_code_block(str(after_data))
        
        self.doc.add_paragraph()
    
    def save(self, output_path):
        """
        Save the document to file.
        
        Args:
            output_path: Path to save DOCX file
        """
        self.doc.save(output_path)
        print(f"✅ Professional document saved: {output_path}")


class MarkdownParser:
    """
    Intelligent Markdown parser that understands document structure.
    
    Parses:
    - Headings (# ## ###)
    - Tables (markdown tables)
    - Bullet points
    - Code blocks (```)
    - JSON blocks
    - Bold/italic text
    - Special sections
    """
    
    def __init__(self, md_file_path):
        """
        Initialize parser with markdown file.
        
        Args:
            md_file_path: Path to markdown file
        """
        with open(md_file_path, 'r', encoding='utf-8') as f:
            self.content = f.read()
        
        self.lines = self.content.split('\n')
    
    def parse(self):
        """
        Parse markdown content into structured elements.
        
        Returns:
            List of tuples: (element_type, element_data)
            
        Element types:
        - 'h1', 'h2', 'h3': Headings
        - 'paragraph': Regular text
        - 'bullet': Bullet point
        - 'code': Code block
        - 'table': Table data
        - 'separator': Horizontal rule
        """
        elements = []
        i = 0
        in_code_block = False
        code_block_lines = []
        in_table = False
        table_lines = []
        
        while i < len(self.lines):
            line = self.lines[i]
            
            # Handle code blocks
            if line.strip().startswith('```'):
                if not in_code_block:
                    in_code_block = True
                    code_block_lines = []
                else:
                    in_code_block = False
                    elements.append(('code', '\n'.join(code_block_lines)))
                    code_block_lines = []
                i += 1
                continue
            
            if in_code_block:
                code_block_lines.append(line)
                i += 1
                continue
            
            # Handle markdown tables
            if '|' in line and line.strip().startswith('|'):
                if not in_table:
                    in_table = True
                    table_lines = [line]
                else:
                    table_lines.append(line)
                i += 1
                
                # Check if table ends
                if i < len(self.lines) and '|' not in self.lines[i]:
                    in_table = False
                    elements.append(('table', self._parse_table(table_lines)))
                    table_lines = []
                continue
            
            # Headings
            if line.startswith('# ') and not line.startswith('## '):
                elements.append(('h1', line[2:].strip()))
            elif line.startswith('## ') and not line.startswith('### '):
                elements.append(('h2', line[3:].strip()))
            elif line.startswith('### '):
                elements.append(('h3', line[4:].strip()))
            
            # Horizontal rules
            elif line.strip() in ['---', '***', '___']:
                elements.append(('separator', None))
            
            # Bullet points (-, *, +, or ✅)
            elif re.match(r'^\s*[-*+✅]\s+', line):
                bullet_text = re.sub(r'^\s*[-*+✅]\s+', '', line).strip()
                is_checked = '✅' in line
                elements.append(('bullet', {'text': bullet_text, 'checked': is_checked}))
            
            # Bold/special text paragraphs
            elif line.strip() and not line.startswith('**'):
                # Check for special prefixes
                if line.startswith('**Problem**:'):
                    elements.append(('problem', line.replace('**Problem**:', '').strip()))
                elif line.startswith('**Solution Implemented**:'):
                    elements.append(('solution', line.replace('**Solution Implemented**:', '').strip()))
                elif line.startswith('**Impact**:'):
                    elements.append(('impact', line.replace('**Impact**:', '').strip()))
                elif line.startswith('**Source File'):
                    elements.append(('metadata', line))
                elif line.startswith('**Output File'):
                    elements.append(('metadata', line))
                elif line.startswith('**Total Records'):
                    elements.append(('metadata', line))
                elif line.strip().startswith('**') and line.strip().endswith('**'):
                    # Section markers
                    elements.append(('section_marker', line.strip('* ')))
                else:
                    elements.append(('paragraph', line))
            
            i += 1
        
        return elements
    
    def _parse_table(self, table_lines):
        """
        Parse markdown table into headers and rows.
        
        Args:
            table_lines: List of table lines from markdown
            
        Returns:
            dict: {'headers': [...], 'rows': [[...]]}
        """
        if len(table_lines) < 2:
            return {'headers': [], 'rows': []}
        
        # Parse headers (first line)
        header_line = table_lines[0]
        headers = [h.strip() for h in header_line.split('|')[1:-1]]
        
        # Skip separator line (second line with dashes)
        # Parse data rows
        rows = []
        for line in table_lines[2:]:
            if line.strip():
                cells = [c.strip() for c in line.split('|')[1:-1]]
                rows.append(cells)
        
        return {'headers': headers, 'rows': rows}


def convert_markdown_to_professional_docx(md_file, output_file, logo_path):
    """
    Main conversion function: Markdown → Professional DOCX
    
    Process:
    1. Parse markdown file intelligently
    2. Create document with professional design
    3. Add logo and branding
    4. Format each element appropriately
    5. Save beautiful DOCX
    
    Args:
        md_file: Path to markdown file
        output_file: Path for output DOCX
        logo_path: Path to MVL logo
    """
    print("=" * 80)
    print("MVL PROFESSIONAL DOCUMENT CONVERTER")
    print("=" * 80)
    print(f"\n📄 Reading markdown file: {md_file}")
    
    # Parse markdown
    parser = MarkdownParser(md_file)
    elements = parser.parse()
    
    print(f"✅ Parsed {len(elements)} document elements")
    print(f"\n🎨 Creating professional design with MVL branding...")
    
    # Create professional document
    designer = ProfessionalDocumentDesigner(logo_path)
    
    # Add title page
    designer.add_title_page(
        title="MVL Supply Chain Intel Hub",
        subtitle="Detailed Data Improvement Documentation",
        date="February 10, 2026",
        version="2.0"
    )
    
    print("✅ Title page created")
    print("📝 Converting content with professional formatting...")
    
    # Process elements
    element_count = {'h1': 0, 'h2': 0, 'h3': 0, 'tables': 0, 'code': 0}
    
    for elem_type, elem_data in elements:
        if elem_type == 'h1':
            designer.add_heading(elem_data, level=1)
            element_count['h1'] += 1
            
        elif elem_type == 'h2':
            designer.add_heading(elem_data, level=2)
            element_count['h2'] += 1
            
        elif elem_type == 'h3':
            designer.add_heading(elem_data, level=3)
            element_count['h3'] += 1
            
        elif elem_type == 'paragraph':
            if elem_data.strip():
                designer.add_paragraph(elem_data)
        
        elif elem_type == 'bullet':
            designer.add_bullet_point(
                elem_data['text'],
                checked=elem_data.get('checked', False)
            )
        
        elif elem_type == 'code':
            designer.add_code_block(elem_data)
            element_count['code'] += 1
        
        elif elem_type == 'table':
            if elem_data['headers'] and elem_data['rows']:
                designer.add_table(elem_data['headers'], elem_data['rows'])
                element_count['tables'] += 1
        
        elif elem_type == 'problem':
            designer.add_info_box("Problem", elem_data, icon="⚠️")
        
        elif elem_type == 'solution':
            designer.add_info_box("Solution Implemented", elem_data, icon="💡")
        
        elif elem_type == 'impact':
            designer.add_info_box("Impact", elem_data, icon="📈")
        
        elif elem_type == 'metadata':
            designer.add_paragraph(elem_data, bold=True, color=designer.COLOR_SECONDARY)
        
        elif elem_type == 'section_marker':
            designer.add_paragraph(elem_data, bold=True, color=designer.COLOR_PRIMARY)
        
        elif elem_type == 'separator':
            designer.doc.add_paragraph()
    
    # Save document
    print(f"\n📊 Document Statistics:")
    print(f"   • Heading 1: {element_count['h1']}")
    print(f"   • Heading 2: {element_count['h2']}")
    print(f"   • Heading 3: {element_count['h3']}")
    print(f"   • Tables: {element_count['tables']}")
    print(f"   • Code Blocks: {element_count['code']}")
    
    print(f"\n💾 Saving professional document...")
    designer.save(output_file)
    
    print(f"\n{'=' * 80}")
    print("✅ CONVERSION COMPLETE!")
    print(f"{'=' * 80}")
    print(f"\n📁 Output file: {output_file}")
    print(f"📏 File size: {os.path.getsize(output_file) / 1024:.1f} KB")
    print(f"\n🎉 Your professional MVL document is ready!")


if __name__ == "__main__":
    # File paths
    BASE_DIR = r"g:\Rita\MVLSupplierIntelHub"
    MD_FILE = os.path.join(BASE_DIR, "DETAILED_IMPROVEMENT_NOTES.md")
    OUTPUT_FILE = os.path.join(BASE_DIR, "MVL_Data_Improvement_Report_Professional.docx")
    LOGO_PATH = os.path.join(BASE_DIR, "MVL Supply Chain Intel Hub - Data", "Logo", "MVLlogo.png")
    
    # Verify files exist
    if not os.path.exists(MD_FILE):
        print(f"❌ Error: Markdown file not found: {MD_FILE}")
        exit(1)
    
    if not os.path.exists(LOGO_PATH):
        print(f"⚠️ Warning: Logo file not found: {LOGO_PATH}")
        print("   Continuing without logo...")
    
    # Convert
    convert_markdown_to_professional_docx(MD_FILE, OUTPUT_FILE, LOGO_PATH)

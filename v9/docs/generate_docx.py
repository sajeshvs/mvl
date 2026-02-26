#!/usr/bin/env python3
"""
MVL Supply Chain Intelligence Hub — DOCX Documentation Generator
Generates a professionally formatted Word document from dashboard analysis.
"""

import os
import json
import io
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ═══════════════════════════════════════════════════════════════
# CONFIGURATION & THEME COLORS
# ═══════════════════════════════════════════════════════════════

THEME = {
    'primary':   RGBColor(0x00, 0x45, 0x78),  # SM Blue
    'secondary': RGBColor(0xD9, 0x6F, 0x3C),  # GSA Orange
    'tertiary':  RGBColor(0x0F, 0x3D, 0x5E),  # M&D Dark Blue
    'accent':    RGBColor(0x2E, 0xCC, 0x71),  # Green
    'text':      RGBColor(0x33, 0x33, 0x33),
    'light':     RGBColor(0x66, 0x66, 0x66),
    'white':     RGBColor(0xFF, 0xFF, 0xFF),
    'bg_light':  RGBColor(0xF5, 0xF7, 0xFA),
}

TAB_COLORS = {
    'SM':  {'hex': '#004578', 'rgb': RGBColor(0x00, 0x45, 0x78), 'name': 'Supplier Marketplace'},
    'GSA': {'hex': '#D96F3C', 'rgb': RGBColor(0xD9, 0x6F, 0x3C), 'name': 'Global Spend Analysis'},
    'MD':  {'hex': '#0F3D5E', 'rgb': RGBColor(0x0F, 0x3D, 0x5E), 'name': 'Materials & Disciplines'},
}

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(os.path.dirname(BASE_DIR), 'data')
OUTPUT_PATH = os.path.join(BASE_DIR, 'MVL_Dashboard_Documentation.docx')

# ═══════════════════════════════════════════════════════════════
# HELPER FUNCTIONS
# ═══════════════════════════════════════════════════════════════

def load_json(filename):
    path = os.path.join(DATA_DIR, filename)
    if os.path.exists(path):
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    return None

def fmt_currency(val):
    if val >= 1e9: return f"${val/1e9:.1f}B"
    if val >= 1e6: return f"${val/1e6:.1f}M"
    if val >= 1e3: return f"${val/1e3:.1f}K"
    return f"${val:,.0f}"

def set_cell_shading(cell, hex_color):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """Set cell border. kwargs: top, bottom, left, right with values like {'sz': '6', 'color': '000000', 'val': 'single'}"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", "4")}" w:space="0" '
            f'w:color="{attrs.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_styled_table(doc, headers, rows, header_color='004578', col_widths=None):
    """Create a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    
    # Style header row
    hdr_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, header_color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Style data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(8.5)
            run.font.color.rgb = THEME['text']
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F5F7FA')
    
    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)
    
    return table

def make_chart_image(fig, dpi=150):
    """Convert matplotlib figure to image bytes."""
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=dpi, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    buf.seek(0)
    plt.close(fig)
    return buf

def add_section_header(doc, text, color_rgb, level=1):
    """Add a colored section header."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = color_rgb

# ═══════════════════════════════════════════════════════════════
# CHART GENERATORS
# ═══════════════════════════════════════════════════════════════

def create_architecture_diagram():
    """Data flow architecture diagram."""
    fig, ax = plt.subplots(figsize=(8, 5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 7)
    ax.axis('off')
    
    # Source data box
    rect = mpatches.FancyBboxPatch((0.5, 5.5), 9, 1.2, boxstyle="round,pad=0.1",
                                    facecolor='#E8F4FD', edgecolor='#004578', linewidth=2)
    ax.add_patch(rect)
    ax.text(5, 6.1, 'SOURCE DATA', ha='center', va='center', fontsize=11, fontweight='bold', color='#004578')
    ax.text(2.5, 5.7, 'PO_List_*.csv\n3,770 rows', ha='center', va='center', fontsize=8, color='#333')
    ax.text(7.5, 5.7, 'Quotation_List_*.xls\n3,921+ records', ha='center', va='center', fontsize=8, color='#333')
    
    # Arrow down
    ax.annotate('', xy=(5, 4.9), xytext=(5, 5.5), arrowprops=dict(arrowstyle='->', color='#666', lw=2))
    
    # Pipeline box
    rect2 = mpatches.FancyBboxPatch((1.5, 3.5), 7, 1.4, boxstyle="round,pad=0.1",
                                     facecolor='#FFF3E6', edgecolor='#D96F3C', linewidth=2)
    ax.add_patch(rect2)
    ax.text(5, 4.5, 'build_v8_data.py (9-Stage Pipeline)', ha='center', va='center',
            fontsize=11, fontweight='bold', color='#D96F3C')
    ax.text(5, 3.85, 'Load → Filter → Dedup → CO Logic → Enrich → SM → GSA → M&D → Save',
            ha='center', va='center', fontsize=7.5, color='#666')
    
    # Arrow down
    ax.annotate('', xy=(5, 2.9), xytext=(5, 3.5), arrowprops=dict(arrowstyle='->', color='#666', lw=2))
    
    # JSON output box
    rect3 = mpatches.FancyBboxPatch((0.5, 1.8), 9, 1.1, boxstyle="round,pad=0.1",
                                     facecolor='#E8F8E8', edgecolor='#2ECC71', linewidth=2)
    ax.add_patch(rect3)
    ax.text(5, 2.5, '7 JSON Files', ha='center', va='center', fontsize=11, fontweight='bold', color='#27ae60')
    ax.text(5, 2.05, 'sm_data  |  gsa_data  |  md_data  |  change_orders  |  conversion_times  |  employees  |  metadata',
            ha='center', va='center', fontsize=7, color='#333')
    
    # Arrow down
    ax.annotate('', xy=(5, 1.2), xytext=(5, 1.8), arrowprops=dict(arrowstyle='->', color='#666', lw=2))
    
    # Browser tabs
    colors_tabs = ['#004578', '#D96F3C', '#0F3D5E']
    labels_tabs = ['SM Tab', 'GSA Tab', 'M&D Tab']
    for i, (c, l) in enumerate(zip(colors_tabs, labels_tabs)):
        x = 2 + i * 3
        rect_t = mpatches.FancyBboxPatch((x - 1, 0.2), 2.4, 0.9, boxstyle="round,pad=0.1",
                                          facecolor=c, edgecolor=c, linewidth=2)
        ax.add_patch(rect_t)
        ax.text(x + 0.2, 0.65, l, ha='center', va='center', fontsize=10, fontweight='bold', color='white')
    
    fig.tight_layout()
    return make_chart_image(fig)

def create_tab_overview_chart():
    """Pie chart showing data distribution across tabs."""
    fig, ax = plt.subplots(figsize=(6, 4))
    labels = ['SM (Quotations)\n3,921 records', 'GSA (Purchase Orders)\n3,746 records',
              'M&D (Combined)\n7,667 records']
    sizes = [3921, 3746, 7667]
    colors = ['#004578', '#D96F3C', '#0F3D5E']
    explode = (0.03, 0.03, 0.03)
    
    wedges, texts, autotexts = ax.pie(sizes, explode=explode, labels=labels, colors=colors,
                                       autopct='%1.1f%%', startangle=90, textprops={'fontsize': 9})
    for t in autotexts:
        t.set_color('white')
        t.set_fontweight('bold')
    ax.set_title('Data Distribution Across Dashboard Tabs', fontsize=12, fontweight='bold', color='#333', pad=15)
    fig.tight_layout()
    return make_chart_image(fig)

def create_co_classification_chart():
    """Stacked bar showing CO classification breakdown."""
    fig, ax = plt.subplots(figsize=(7, 3.5))
    categories = ['Base PO\n(Rev 1)', 'Change Order\n(Rev 2-6)', 'Standalone\n(Rev 7+)', 'Other Prefix\n(SPO, etc.)']
    counts = [3342, 296, 12, 96]
    colors = ['#2ECC71', '#E74C3C', '#3498DB', '#95A5A6']
    
    bars = ax.barh(categories, counts, color=colors, height=0.6, edgecolor='white', linewidth=1)
    for bar, count in zip(bars, counts):
        ax.text(bar.get_width() + 30, bar.get_y() + bar.get_height()/2,
                f'{count:,}', va='center', ha='left', fontsize=10, fontweight='bold', color='#333')
    
    ax.set_xlabel('Number of POs', fontsize=10, color='#666')
    ax.set_title('PO Classification Breakdown (3,746 Total)', fontsize=12, fontweight='bold', color='#333')
    ax.set_xlim(0, max(counts) * 1.15)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    fig.tight_layout()
    return make_chart_image(fig)

def create_currency_chart():
    """Bar chart showing FX rates."""
    fig, ax = plt.subplots(figsize=(7, 3))
    currencies = ['AED', 'SAR', 'EUR', 'GBP', 'INR', 'KWD', 'QAR', 'BHD', 'OMR', 'ZAR', 'SGD', 'PKR']
    rates = [3.6725, 3.75, 0.92, 0.79, 83.0, 0.3077, 3.64, 0.376, 0.385, 18.5, 1.34, 278.0]
    
    colors = ['#004578' if r < 5 else '#D96F3C' for r in rates]
    bars = ax.bar(currencies, rates, color=colors, edgecolor='white', linewidth=0.5)
    
    for bar, rate in zip(bars, rates):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + max(rates)*0.02,
                f'{rate}', ha='center', va='bottom', fontsize=7, color='#666')
    
    ax.set_ylabel('Units per 1 USD', fontsize=9, color='#666')
    ax.set_title('Exchange Rates (19 Currencies Supported)', fontsize=11, fontweight='bold', color='#333')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.xticks(rotation=45, fontsize=8)
    fig.tight_layout()
    return make_chart_image(fig)

def create_pipeline_stages_chart():
    """Horizontal bar showing pipeline stages."""
    fig, ax = plt.subplots(figsize=(7, 4))
    stages = ['[9/9] Save JSON', '[8/9] Build M&D', '[7/9] Build GSA',
              '[6/9] Build SM', '[5/9] Enrich POs', '[4/9] Calc CO',
              '[3/9] Dedup POs', '[2/9] Filter Quotes', '[1/9] Load Data']
    
    colors = ['#95A5A6', '#0F3D5E', '#D96F3C', '#004578', '#27AE60',
              '#E74C3C', '#3498DB', '#F39C12', '#2C3E50']
    
    widths = [1, 3, 4, 3, 2, 2, 1.5, 1.5, 2]
    
    bars = ax.barh(stages, widths, color=colors, height=0.65, edgecolor='white', linewidth=1)
    
    labels = ['7 JSON files', '7,667 combined records', '3,746 POs + rankings',
              '3,921 quotes + stats', '180 Q→PO links', '296 COs, 192 groups',
              '3,746 unique POs', '3,921 clean RFQs', 'CSV + XLS auto-detect']
    
    for bar, lbl in zip(bars, labels):
        ax.text(bar.get_width() + 0.1, bar.get_y() + bar.get_height()/2,
                lbl, va='center', ha='left', fontsize=8, color='#666')
    
    ax.set_xlim(0, 8)
    ax.set_title('Data Pipeline — 9 Processing Stages', fontsize=12, fontweight='bold', color='#333')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.xaxis.set_visible(False)
    fig.tight_layout()
    return make_chart_image(fig)

def create_sm_kpi_cards_chart():
    """Visual card layout for SM KPIs."""
    fig, axes = plt.subplots(1, 4, figsize=(8, 2))
    kpis = [
        ('RFQ Count', '3,921', '#004578'),
        ('Quote Value', '$478.6M', '#004578'),
        ('Total POs', '3,746', '#004578'),
        ('Win Rate', '95.3%', '#2ECC71'),
    ]
    for ax, (title, val, color) in zip(axes, kpis):
        ax.set_xlim(0, 1)
        ax.set_ylim(0, 1)
        ax.axis('off')
        rect = mpatches.FancyBboxPatch((0.05, 0.05), 0.9, 0.9, boxstyle="round,pad=0.05",
                                        facecolor=color, edgecolor='white', linewidth=2)
        ax.add_patch(rect)
        ax.text(0.5, 0.62, val, ha='center', va='center', fontsize=16, fontweight='bold', color='white')
        ax.text(0.5, 0.3, title, ha='center', va='center', fontsize=9, color='#FFFFFFCC')
    fig.tight_layout(pad=0.5)
    return make_chart_image(fig)

def create_gsa_kpi_cards_chart():
    """Visual card layout for GSA KPIs."""
    fig, axes = plt.subplots(1, 4, figsize=(8, 2))
    kpis = [
        ('Total POs', '3,746', '#D96F3C'),
        ('Total Spend', '$481.1M', '#27AE60'),
        ('Change Orders', '296', '#3498DB'),
        ('CO Value', '$12.1M', '#F39C12'),
    ]
    for ax, (title, val, color) in zip(axes, kpis):
        ax.set_xlim(0, 1)
        ax.set_ylim(0, 1)
        ax.axis('off')
        rect = mpatches.FancyBboxPatch((0.05, 0.05), 0.9, 0.9, boxstyle="round,pad=0.05",
                                        facecolor=color, edgecolor='white', linewidth=2)
        ax.add_patch(rect)
        ax.text(0.5, 0.62, val, ha='center', va='center', fontsize=16, fontweight='bold', color='white')
        ax.text(0.5, 0.3, title, ha='center', va='center', fontsize=9, color='#FFFFFFCC')
    fig.tight_layout(pad=0.5)
    return make_chart_image(fig)

def create_material_codes_chart():
    """Horizontal bar for material code classification."""
    fig, ax = plt.subplots(figsize=(7, 4))
    codes = ['Architectural', 'Chemicals', 'Electrical', 'Fire', 'Logistics',
             'Mechanical', 'Office Assets', 'Protection', 'Rental', 'Services', 'Tools', 'Various']
    raw_counts = [8, 2, 1, 7, 4, 2, 1, 1, 1, 5, 1, 5]  # number of raw materials per category
    
    colors = ['#2B4257', '#3B82F6', '#60A5FA', '#EF4444', '#F59E0B',
              '#06B6D4', '#10B981', '#8B5CF6', '#22C55E', '#0F3D5E', '#9333EA', '#64748B']
    
    bars = ax.barh(codes, raw_counts, color=colors, height=0.6, edgecolor='white')
    for bar, cnt in zip(bars, raw_counts):
        ax.text(bar.get_width() + 0.1, bar.get_y() + bar.get_height()/2,
                f'{cnt} materials', va='center', fontsize=8, color='#666')
    
    ax.set_xlabel('Number of Raw Materials', fontsize=9, color='#666')
    ax.set_title('12 Material Code Categories & Raw Material Counts', fontsize=11, fontweight='bold', color='#333')
    ax.set_xlim(0, max(raw_counts) + 2)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    fig.tight_layout()
    return make_chart_image(fig)

def create_status_breakdown_chart():
    """Doughnut chart for SM status distribution."""
    fig, ax = plt.subplots(figsize=(5, 4))
    statuses = ['Order', 'Quotation', 'Waiting', 'Cancelled', 'Closed']
    colors = ['#4CAF50', '#2196F3', '#FFC107', '#F44336', '#9E9E9E']
    # Approximate distribution
    sizes = [55, 25, 10, 5, 5]
    wedges, texts, autotexts = ax.pie(sizes, labels=statuses, colors=colors,
                                       autopct='%1.0f%%', startangle=90,
                                       pctdistance=0.75, textprops={'fontsize': 9})
    for t in autotexts:
        t.set_fontweight('bold')
    centre_circle = plt.Circle((0, 0), 0.50, fc='white')
    ax.add_artist(centre_circle)
    ax.set_title('Quotation Status Distribution', fontsize=11, fontweight='bold', color='#333')
    fig.tight_layout()
    return make_chart_image(fig)


# ═══════════════════════════════════════════════════════════════
# DOCUMENT BUILDER
# ═══════════════════════════════════════════════════════════════

def build_document():
    doc = Document()
    
    # ── Page Setup ──
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
    
    # ── Define Styles ──
    style = doc.styles['Normal']
    style.font.name = 'Segoe UI'
    style.font.size = Pt(10)
    style.font.color.rgb = THEME['text']
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    
    for level in range(1, 5):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Segoe UI'
        h_style.font.color.rgb = THEME['primary']
        if level == 1:
            h_style.font.size = Pt(20)
            h_style.paragraph_format.space_before = Pt(24)
        elif level == 2:
            h_style.font.size = Pt(16)
            h_style.paragraph_format.space_before = Pt(18)
        elif level == 3:
            h_style.font.size = Pt(13)
            h_style.paragraph_format.space_before = Pt(12)
        else:
            h_style.font.size = Pt(11)
    
    # ════════════════════════════════════════════════════════════
    # COVER PAGE
    # ════════════════════════════════════════════════════════════
    
    for _ in range(6):
        doc.add_paragraph('')
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('MVL Supply Chain\nIntelligence Hub')
    run.font.size = Pt(36)
    run.font.color.rgb = THEME['primary']
    run.bold = True
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('V8 Dashboard — Technical Documentation')
    run.font.size = Pt(18)
    run.font.color.rgb = THEME['secondary']
    
    doc.add_paragraph('')
    
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run('━' * 50)
    run.font.color.rgb = THEME['light']
    run.font.size = Pt(10)
    
    doc.add_paragraph('')
    
    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = details.add_run('Data Sources  •  Calculation Logic  •  Display Reference\n\n')
    run.font.size = Pt(12)
    run.font.color.rgb = THEME['light']
    
    run = details.add_run('February 2026  |  Version 8.0\n')
    run.font.size = Pt(11)
    run.font.color.rgb = THEME['text']
    
    run = details.add_run('https://sajeshvs.github.io/mvl/v8/')
    run.font.size = Pt(10)
    run.font.color.rgb = THEME['primary']
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ════════════════════════════════════════════════════════════
    
    add_section_header(doc, 'Table of Contents', THEME['primary'], level=1)
    
    toc_items = [
        ('1', 'Executive Summary'),
        ('2', 'Architecture Overview'),
        ('3', 'Data Pipeline (9 Stages)'),
        ('4', 'Currency Conversion'),
        ('5', 'Change Order Classification Logic'),
        ('6', 'Tab 1 — Supplier Marketplace (SM)'),
        ('7', 'Tab 2 — Global Spend Analysis (GSA)'),
        ('8', 'Tab 3 — Materials & Disciplines (M&D)'),
        ('9', 'Data File Reference'),
        ('10', 'Material Code Classification'),
        ('11', 'Country Normalization'),
    ]
    
    for num, title_text in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{num}.  ')
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = THEME['primary']
        run = p.add_run(title_text)
        run.font.size = Pt(11)
        run.font.color.rgb = THEME['text']
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ════════════════════════════════════════════════════════════
    
    add_section_header(doc, '1. Executive Summary', THEME['primary'], level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('The MVL Supply Chain Intelligence Hub ')
    run.bold = True
    run = p.add_run(
        'is a browser-based, three-tab analytics dashboard that provides end-to-end visibility '
        'into procurement operations — from initial Request for Quotation (RFQ) through Purchase '
        'Order (PO) execution and Change Order (CO) tracking.'
    )
    
    add_section_header(doc, 'Dashboard Tabs at a Glance', THEME['primary'], level=2)
    
    add_styled_table(doc,
        ['Tab', 'Theme', 'Focus Area', 'Primary Data'],
        [
            ['Supplier Marketplace', 'Blue #004578', 'RFQ pipeline, quotation tracking, supplier discovery', 'Quotation records (RFQs)'],
            ['Global Spend Analysis', 'Orange #D96F3C', 'PO spend, change orders, entity & project analysis', 'Purchase Orders'],
            ['Materials & Disciplines', 'Dark Blue #0F3D5E', 'Material categorization, discipline spend, conversions', 'Combined RFQs + POs'],
        ],
        header_color='004578'
    )
    
    doc.add_paragraph('')
    
    # Tab overview chart
    chart_buf = create_tab_overview_chart()
    doc.add_picture(chart_buf, width=Inches(4.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_section_header(doc, 'Key Metrics Summary', THEME['primary'], level=2)
    
    add_styled_table(doc,
        ['Metric', 'Value', 'Source'],
        [
            ['Total RFQ Records', '3,921', 'Quotation Excel export'],
            ['Total Purchase Orders', '3,746', 'PO CSV export'],
            ['Total PO Spend (USD)', '~$481M', 'Converted via FX rates'],
            ['Change Orders', '296 (192 groups)', 'PO/RFPO rev 2–6'],
            ['Master Supplier Count', '2,189', 'suppliers.json'],
            ['Active Suppliers (in POs)', '1,133', 'Derived from PO data'],
            ['Material Categories', '12', 'Code-based classification'],
            ['Raw Materials', '33', 'Distinct material names'],
        ],
        header_color='004578'
    )
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════════════
    # 2. ARCHITECTURE OVERVIEW
    # ════════════════════════════════════════════════════════════
    
    add_section_header(doc, '2. Architecture Overview', THEME['primary'], level=1)
    
    add_section_header(doc, 'Technology Stack', THEME['primary'], level=2)
    
    add_styled_table(doc,
        ['Layer', 'Technology', 'Purpose'],
        [
            ['Frontend', 'HTML5, CSS3, Vanilla JavaScript', 'Single-page application'],
            ['Charts', 'Chart.js 4.x', 'Bar, line, doughnut, radar charts'],
            ['Maps', 'Leaflet.js 1.9.4', 'Interactive supplier location map'],
            ['Data Pipeline', 'Python 3.12', 'Excel/CSV → JSON transformation'],
            ['Hosting', 'GitHub Pages', 'Static file deployment'],
        ],
        header_color='004578'
    )
    
    doc.add_paragraph('')
    
    add_section_header(doc, 'Data Flow Architecture', THEME['primary'], level=2)
    
    chart_buf = create_architecture_diagram()
    doc.add_picture(chart_buf, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    
    add_section_header(doc, 'File Structure', THEME['primary'], level=2)
    
    add_styled_table(doc,
        ['File', 'Purpose'],
        [
            ['v8/index.html', 'Single-page app with 3 tabs (~1,063 lines)'],
            ['v8/shared/scripts.js', 'All dashboard logic (~5,928 lines)'],
            ['v8/shared/styles.css', 'Complete CSS with design tokens (~2,820 lines)'],
            ['v8/data/build_v8_data.py', 'Python pipeline — 9 stages (~1,508 lines)'],
            ['v8/data/sm_data.json', 'Supplier Marketplace data'],
            ['v8/data/gsa_data.json', 'Global Spend Analysis data'],
            ['v8/data/md_data.json', 'Materials & Disciplines data'],
            ['v8/data/suppliers.json', 'Master supplier directory (2,189 entries)'],
            ['v8/data/change_orders.json', 'CO group details (192 groups)'],
            ['v8/data/conversion_times.json', 'RFQ→PO linkage & timing (180 links)'],
            ['v8/data/client_country_map.json', 'Client→Country mapping (1,098 entries)'],
            ['v8/data/entity_code_map.json', 'Entity code→name lookup (~20 entries)'],
        ],
        header_color='004578'
    )
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════════════
    # 3. DATA PIPELINE
    # ════════════════════════════════════════════════════════════
    
    add_section_header(doc, '3. Data Pipeline (9 Stages)', THEME['secondary'], level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('The Python pipeline (build_v8_data.py) is the single source of truth ')
    run.bold = True
    run = p.add_run('for all dashboard data. It reads raw Excel/CSV files and transforms them '
                     'into optimized JSON through 9 sequential stages.')
    
    doc.add_paragraph('')
    chart_buf = create_pipeline_stages_chart()
    doc.add_picture(chart_buf, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    
    stages_data = [
        ['[1/9]', 'Load Data Files', 'Auto-detects CSV from Data-New/ folder; falls back to legacy XLS via xlrd', '3,770 PO rows + 3,921+ quotation rows'],
        ['[2/9]', 'Filter Quotations', 'Keeps RFQ only (removes Internal Quotations), deduplicates, removes empties', '3,921 clean quotation records'],
        ['[3/9]', 'Deduplicate POs', 'Removes 24 duplicate PO numbers, converts all values to USD', '3,746 unique POs'],
        ['[4/9]', 'Calculate Change Orders', '3-tier CO logic: PO/RFPO rev 2–6 = CO, rev 7+ = standalone', '296 COs in 192 groups'],
        ['[5/9]', 'Enrich POs', 'Matches POs to quotations by orderId — fills material, project, entity', '180 Q→PO links, avg 29.5 days'],
        ['[6/9]', 'Build SM Data', 'Quotation stats, status counts, employee records, entity aggregations', 'sm_data.json'],
        ['[7/9]', 'Build GSA Data', 'PO spend summaries, supplier rankings, CO deduction logic, trends', 'gsa_data.json'],
        ['[8/9]', 'Build M&D Data', 'Combined RFQ + PO by material code, discipline breakdown', 'md_data.json'],
        ['[9/9]', 'Save All JSON', 'Conversion times, change order groups, metadata, all files saved', '7 JSON output files'],
    ]
    
    add_styled_table(doc,
        ['Stage', 'Name', 'Description', 'Output'],
        stages_data,
        header_color='D96F3C'
    )
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════════════
    # 4. CURRENCY CONVERSION
    # ════════════════════════════════════════════════════════════
    
    add_section_header(doc, '4. Currency Conversion', THEME['primary'], level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('Formula: ')
    run.bold = True
    run = p.add_run('USD Value = Original Amount ÷ Exchange Rate')
    run.italic = True
    
    doc.add_paragraph('')
    chart_buf = create_currency_chart()
    doc.add_picture(chart_buf, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    
    fx_data = [
        ['USD', '1.0000', 'Base currency'],
        ['AED', '3.6725', 'UAE Dirham'],
        ['SAR', '3.7500', 'Saudi Riyal'],
        ['EUR / EURO', '0.9200', 'Euro (both codes mapped)'],
        ['GBP', '0.7900', 'British Pound'],
        ['INR', '83.0000', 'Indian Rupee'],
        ['KWD', '0.3077', 'Kuwaiti Dinar'],
        ['QAR', '3.6400', 'Qatari Riyal'],
        ['BHD', '0.3760', 'Bahraini Dinar'],
        ['OMR', '0.3850', 'Omani Rial'],
        ['NPR', '133.5000', 'Nepalese Rupee (PO override: 1:1)'],
        ['JPY', '149.5000', 'Japanese Yen (PO override: 1:1)'],
        ['ZAR', '18.5000', 'South African Rand'],
        ['SGD', '1.3400', 'Singapore Dollar'],
        ['PKR', '278.0000', 'Pakistani Rupee'],
        ['EGP', '30.9000', 'Egyptian Pound'],
        ['JOD', '0.7090', 'Jordanian Dinar'],
        ['LKR', '320.0000', 'Sri Lankan Rupee'],
    ]
    
    add_styled_table(doc,
        ['Currency', 'Rate (per 1 USD)', 'Notes'],
        fx_data,
        header_color='004578'
    )
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run('PO-specific overrides: ')
    run.bold = True
    run = p.add_run('NPR and JPY are treated as 1:1 with USD for PO values — these appear to be data-entry '
                     'artifacts where USD values were entered under NPR/JPY currencies.')
    run.font.size = Pt(9)
    run.font.color.rgb = THEME['light']
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════════════
    # 5. CHANGE ORDER CLASSIFICATION
    # ════════════════════════════════════════════════════════════
    
    add_section_header(doc, '5. Change Order Classification Logic', THEME['primary'], level=1)
    
    add_section_header(doc, 'PO Number Structure', THEME['primary'], level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('PO-1234-M4004-3')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Consolas'
    
    p = doc.add_paragraph()
    run = p.add_run('  PO      = Prefix (PO, RFPO, SPO, RFSPO, etc.)\n')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run = p.add_run('  1234    = Sequential Number\n')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run = p.add_run('  M4004   = Entity Code (maps to entity name)\n')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run = p.add_run('  3       = Revision Number')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    add_section_header(doc, '3-Tier Classification Rules', THEME['primary'], level=2)
    
    doc.add_paragraph('')
    chart_buf = create_co_classification_chart()
    doc.add_picture(chart_buf, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    
    add_styled_table(doc,
        ['Tier', 'Prefix', 'Revision', 'Classification', 'Grouping'],
        [
            ['1', 'PO / RFPO', '1', 'Base PO', 'Shared baseGroupKey'],
            ['2', 'PO / RFPO', '2–6', 'Change Order', 'Shared baseGroupKey with rev 1'],
            ['3', 'PO / RFPO', '7+', 'Independent Standalone', 'Own baseGroupKey (self)'],
            ['4', 'All others', 'Any', 'Standalone', 'Own baseGroupKey (self)'],
        ],
        header_color='004578'
    )
    
    add_section_header(doc, 'CO Value Calculation (Incremental Deduction)', THEME['primary'], level=2)
    
    p = doc.add_paragraph('Change Order values represent the incremental cost difference, not the absolute PO value:')
    
    add_styled_table(doc,
        ['Scenario', 'CO Value Formula'],
        [
            ['First CO in group (or orphan)', 'Full CO amount'],
            ['Same value as previous version', '$0 (count only)'],
            ['Non-consecutive revision gap', 'Full CO amount'],
            ['Normal consecutive revision', 'CO amount − Previous version amount'],
        ],
        header_color='E74C3C'
    )
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════════════
    # 6. SM TAB
    # ════════════════════════════════════════════════════════════
    
    add_section_header(doc, '6. Tab 1 — Supplier Marketplace (SM)', TAB_COLORS['SM']['rgb'], level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('Theme: Blue (#004578)  |  Focus: RFQ pipeline and supplier discovery')
    run.font.color.rgb = THEME['light']
    run.italic = True
    
    # SM Data Sources
    add_section_header(doc, '6.1 Data Sources', TAB_COLORS['SM']['rgb'], level=2)
    
    add_styled_table(doc,
        ['File', 'Variable', 'Content'],
        [
            ['sm_data.json', 'smData', 'Quotation workbench, status/entity/material summaries, MVL employee records, filter lists'],
            ['gsa_data.json', 'gsaData', 'PO counts, CO counts/values, supplier rankings, entity spend (for PO-related KPIs)'],
            ['suppliers.json', 'suppliersData', 'Master supplier directory — name, contact, email, phone, rating, address, country'],
            ['conversion_times.json', '_conversionTimes', 'RFQ→PO monthly conversion averages'],
            ['client_country_map.json', 'clientCountryMap', 'Client name → country mapping (1,098 entries)'],
        ],
        header_color='004578'
    )
    
    # SM Filters
    add_section_header(doc, '6.2 Filters (10 Controls)', TAB_COLORS['SM']['rgb'], level=2)
    
    add_styled_table(doc,
        ['#', 'Filter', 'HTML ID', 'Options Source', 'Matches Field'],
        [
            ['1', 'Entity', 'filterEntity', 'smData.entities[].Entity', 'q.Entity'],
            ['2', 'Project', 'filterProject', 'smData.workbench[] (2+ quotations)', 'q.ProjectName'],
            ['3', 'Supplier', 'filterSupplier', 'gsaData.filters.suppliers', 'q.Client'],
            ['4', 'Status', 'filterStatus', 'Hardcoded: Order, Quotation, Waiting, Cancelled, Closed', 'q.Status'],
            ['5', 'Material', 'filterMaterial', 'smData.filters.materials (30 names)', 'q.Material'],
            ['6', 'Material Code', 'filterMaterialCode', 'smData.filters.materialCodes (12 categories)', 'q.materialCode'],
            ['7', 'Date From', 'filterDateFrom', 'User-selected', 'q.Date >= fromDate'],
            ['8', 'Date To', 'filterDateTo', 'User-selected', 'q.Date <= toDate'],
            ['9', 'Search', 'searchInput', 'Free text', 'QuotationNumber, Entity, ProjectName, Description, Client'],
            ['10', 'Clear All', 'Button', '—', 'Resets all filters'],
        ],
        header_color='004578'
    )
    
    p = doc.add_paragraph()
    run = p.add_run('Filter Pipeline: ')
    run.bold = True
    run = p.add_run('applyFilters() filters smData.workbench[] array, then updates ALL visual components from the filtered result set.')
    run.font.size = Pt(9)
    
    # SM KPIs
    add_section_header(doc, '6.3 KPI Cards (7 Metrics)', TAB_COLORS['SM']['rgb'], level=2)
    
    chart_buf = create_sm_kpi_cards_chart()
    doc.add_picture(chart_buf, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    
    add_styled_table(doc,
        ['#', 'KPI Title', 'Formula', 'Refilters?'],
        [
            ['1', 'Request for Quotation', 'COUNT(smData.workbench)', 'Yes'],
            ['2', 'Quote Value', 'SUM(convertToUSD(q.QuotationValue, q.Currency))', 'Yes'],
            ['3', 'Total Purchase Orders', 'COUNT(non_spo_pos) from pipeline', 'No — pre-calculated'],
            ['4', 'PO Values', 'SUM(po.poSpendUSD) for non-SPO POs', 'No — pre-calculated'],
            ['5', 'Win Rate', 'totalPOs / totalQuotations × 100', 'Yes'],
            ['6', 'Change Orders', 'COUNT(gsaData WHERE poType="Change Order")', 'No — from GSA data'],
            ['7', 'CO Value', 'Incremental deduction logic (Section 5)', 'No — from GSA data'],
        ],
        header_color='004578'
    )
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run('Note: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
    run = p.add_run('KPIs 3, 4, 6, 7 are sourced from PO/GSA data and remain constant regardless of SM quotation filters. '
                     'Only KPIs 1, 2, 5 respond to filter changes.')
    run.font.size = Pt(9)
    run.font.color.rgb = THEME['light']
    
    # SM Charts
    add_section_header(doc, '6.4 Charts and Visualizations', TAB_COLORS['SM']['rgb'], level=2)
    
    sm_charts = [
        ['Status Chart', 'Custom HTML bar list', 'statusChart', 'smData.statusSummary[] — {Status, Count, TotalValueUSD}',
         'Order=#4CAF50, Quotation=#2196F3, Waiting=#FFC107, Cancelled=#F44336, Closed=#9E9E9E',
         'Click bar → filters by that status. Sub-KPIs: Conversion Rate, Open Quotes (Quotation + Waiting)'],
        ['Entity Comparison', 'Chart.js horizontal bar', 'entityChartCanvas', 'quote view: smData.entities[].TotalValueUSD; spend view: gsaData.entityBreakdown[].valueUSD',
         'Toggle: "By Quote" / "By PO Spend". Frozen x-axis, scrollable.',
         'Click bar → filters by entity. Dynamic 28px per entity bar height.'],
        ['Top 10 Suppliers', 'Custom HTML ranked list', 'topSuppliers', 'gsaData.supplierRankings.top[] (first 10)',
         'Gold/silver/bronze rank circles, proportional bar widths',
         'Click row → updates Supplier Profile + sets Supplier filter'],
        ['Supplier Map', 'Leaflet.js interactive map', 'supplierMap', 'suppliersData.suppliers[] grouped by country via normalizeCountry()',
         'CircleMarker radius 8–25px, 5-color intensity scale (green→red)',
         'Popup: country, count, value, top 5 clients. Uses clientCountryMap + entityCountryMap fallback.'],
        ['Material Distribution', 'Chart.js (Bar/Pie/Line/Radar)', 'materialChartCanvas', 'smData.materialsByDiscipline[] (top 8 by value)',
         '4-way toggle: Bar, Pie, Line, Radar views',
         'Click segment → filters by material'],
        ['Quotation to PO Time', 'Chart.js vertical bar', 'quotationTimeChart', 'conversion_times.json → monthlyAverage[]',
         'Labels show "Nd" (days) above bars. Tooltip: X days (Y POs)',
         'Date-range filtered by month (YYYY-MM comparison)'],
        ['Submit & Order Trend', 'Chart.js multi-line (3 datasets)', 'trendChart', 'Quotes from submission_date, Orders from po_date, COs from GSA workbench',
         'Quotes (#0066CC), Orders (#339933), COs (#FF9900)',
         'NOT refiltered — rendered once at initial load'],
    ]
    
    add_styled_table(doc,
        ['Chart Name', 'Type', 'Canvas/Container', 'Data Source', 'Styling', 'Interaction & Notes'],
        sm_charts,
        header_color='004578'
    )
    
    # SM Supplier Profile
    add_section_header(doc, '6.5 Supplier Profile Card', TAB_COLORS['SM']['rgb'], level=2)
    
    add_styled_table(doc,
        ['Field', 'HTML ID', 'Data Source'],
        [
            ['Avatar', 'supplierAvatar', 'First character of supplier name'],
            ['Name', 'supplierName', 'Selected supplier name'],
            ['Location', 'supplierLocation', 'suppliersData → normalizeCountry(address.country_standardized)'],
            ['Contact', 'supplierContact', 'suppliersData → contact.primary_contact'],
            ['Email', 'supplierEmail', 'suppliersData → contact.email'],
            ['Phone', 'supplierPhone', 'suppliersData → contact.phone'],
            ['Rating', 'supplierRating', 'suppliersData → rating.score → rendered as ★/☆'],
        ],
        header_color='004578'
    )
    
    # SM Employee List
    add_section_header(doc, '6.6 Responsible MVL Employee', TAB_COLORS['SM']['rgb'], level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('Data source: ')
    run.bold = True
    run = p.add_run('smData.suppliers[] — these are MVL procurement contacts (e.g., "Lince M.", "Marman I."), '
                     'NOT vendor companies. Ranked list with gold/silver/bronze circles, togglable "By Spend" / "By Count" sort.')
    
    # SM Bottom Tables
    add_section_header(doc, '6.7 Bottom Tables', TAB_COLORS['SM']['rgb'], level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('Quotation Details Table ')
    run.bold = True
    run = p.add_run('(default tab) — ')
    run = p.add_run('Columns: Quotation Number, Status (color badge), Material, Project, Value (USD), Contact. '
                     'Pagination: 25/50/100/200 rows per page.')
    
    p = doc.add_paragraph()
    run = p.add_run('Supplier List Table ')
    run.bold = True
    run = p.add_run('— Columns: Supplier Name, Contact, Email, Phone, Country, Category. '
                     'Source: suppliers.json (2,189 suppliers). Click row → updates Supplier Profile.')
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════════════
    # 7. GSA TAB
    # ════════════════════════════════════════════════════════════
    
    add_section_header(doc, '7. Tab 2 — Global Spend Analysis (GSA)', TAB_COLORS['GSA']['rgb'], level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('Theme: Orange (#D96F3C)  |  Focus: PO spend analysis and change order tracking')
    run.font.color.rgb = THEME['light']
    run.italic = True
    
    # GSA Data Sources
    add_section_header(doc, '7.1 Data Sources', TAB_COLORS['GSA']['rgb'], level=2)
    
    add_styled_table(doc,
        ['File', 'Variable', 'Content'],
        [
            ['gsa_data.json', 'gsaData', 'PO workbench (3,746 records), spend summary, supplier rankings, entity/material breakdown, monthly trends, CO details'],
            ['suppliers.json', 'suppliersData', 'Supplier details for supplier card (name, contact, rating, location)'],
        ],
        header_color='D96F3C'
    )
    
    # GSA Filters
    add_section_header(doc, '7.2 Filters (10 Controls)', TAB_COLORS['GSA']['rgb'], level=2)
    
    add_styled_table(doc,
        ['#', 'Filter', 'HTML ID', 'Options Source', 'Matches Field'],
        [
            ['1', 'Entity', 'gsaFilterEntity', 'gsaData.filters.entities', 'po.entity'],
            ['2', 'Supplier', 'gsaFilterSupplier', 'gsaData.filters.suppliers', 'po.supplier'],
            ['3', 'Project', 'gsaFilterProject', 'Derived unique po.project values', 'po.project'],
            ['4', 'Material', 'gsaFilterMaterial', 'gsaData.filters.materials', 'po.material'],
            ['5', 'Material Code', 'gsaFilterMaterialCode', 'gsaData.filters.materialCodes', 'po.materialCode'],
            ['6', 'PO Type', 'gsaFilterDiscipline', 'gsaData.filters.poTypes', 'po.poType'],
            ['7', 'Year', 'gsaFilterYear', 'gsaData.filters.years', 'po.year'],
            ['8', 'Date From', 'gsaFilterFrom', 'Date input', 'po.poDate >= fromDate'],
            ['9', 'Date To', 'gsaFilterTo', 'Date input', 'po.poDate <= toDate'],
            ['10', 'Search', 'gsaSearchInput', 'Free text (debounced 300ms)', 'poNumber, poName, project, supplier, material, entity, orderId'],
        ],
        header_color='D96F3C'
    )
    
    # GSA KPIs
    add_section_header(doc, '7.3 KPI Cards (6 Metrics)', TAB_COLORS['GSA']['rgb'], level=2)
    
    chart_buf = create_gsa_kpi_cards_chart()
    doc.add_picture(chart_buf, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    
    add_styled_table(doc,
        ['#', 'KPI Title', 'Unfiltered Source', 'Filtered Calculation'],
        [
            ['1', 'Total No. of Purchase Orders', 'gsaData.summary.totalPOs', 'filteredPOs.length'],
            ['2', 'Total Spend', 'gsaData.summary.totalSpendUSD', 'SUM(convertToUSD(po.valueUSD, po.currency))'],
            ['3', 'Total No. of Change Orders', 'gsaData.summary.changeOrders', 'COUNT(po WHERE poType="Change Order")'],
            ['4', 'Total Amount of Change Orders', 'gsaData.summary.changeOrderValue', 'SUM(convertToUSD(co.valueUSD, co.currency))'],
            ['5', 'No. of Suppliers', 'gsaData.summary.supplierCount (2,189)', 'COUNT(DISTINCT po.supplier)'],
            ['6', 'No. of Entities', 'gsaData.summary.entityCount', 'COUNT(DISTINCT po.entity)'],
        ],
        header_color='D96F3C'
    )
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run('Sub-labels: ')
    run.bold = True
    run = p.add_run('KPI 3 shows "N groups" count. KPI 4 shows "X% of total spend" = (CO Value / Total Spend) × 100.')
    run.font.size = Pt(9)
    run.font.color.rgb = THEME['light']
    
    # GSA Charts
    add_section_header(doc, '7.4 Charts and Visualizations', TAB_COLORS['GSA']['rgb'], level=2)
    
    gsa_charts = [
        ['Annual Spend Trend', 'Stacked bar + line combo', 'gsaSpendTrendChart',
         'Base Spend (bar, orange #FF8C00) + Change Orders (bar, gold #FFD700) + Running Total (line, blue #0066CC)',
         'Unfiltered: gsaData.monthlyTrend[]. Filtered: grouped from filteredData by po.yearMonth'],
        ['Spend by Entity', 'Horizontal bar', 'gsaEntityChart',
         'Top 8 Entities by PO Value. Click → cross-filters to that entity',
         'Unfiltered: gsaData.entityBreakdown[]. Filtered: grouped by po.entity'],
        ['Spend by Projects', 'Horizontal bar', 'gsaProjectChart',
         'Top 8 Projects by PO Value. Click → cross-filters to that project',
         'Always calculated from filteredData by po.project, top 8'],
        ['Top Suppliers', 'Horizontal bar', 'gsaTopSuppliersChart',
         'Top 10 Suppliers by Spend. Click → updates supplier card + cross-filters',
         'Unfiltered: gsaData.supplierRankings.top[]. Filtered: grouped by po.supplier'],
        ['Most Inactive Suppliers', 'Horizontal bar', 'gsaBottomSuppliersChart',
         'Bottom 10 Suppliers by Spend. Click → updates supplier card + cross-filters',
         'Same aggregation, sorted ascending, bottom 10'],
    ]
    
    add_styled_table(doc,
        ['Chart Name', 'Type', 'Canvas ID', 'Datasets & Interaction', 'Data Source'],
        gsa_charts,
        header_color='D96F3C'
    )
    
    # GSA Supplier Card
    add_section_header(doc, '7.5 Supplier Details Card', TAB_COLORS['GSA']['rgb'], level=2)
    
    add_styled_table(doc,
        ['Field', 'HTML ID', 'Data Source'],
        [
            ['Name', 'gsaSupplierName', 'Selected supplier name'],
            ['Location', 'gsaSupplierLocation', 'suppliersData → normalizeCountry()'],
            ['Stars', 'gsaSupplierStars', 'suppliersData → rating.score'],
            ['Rating', 'gsaSupplierRating', 'rating.toFixed(2) + "/5"'],
            ['Email', 'gsaSupplierEmail', 'suppliersData → contact.email'],
            ['Contact', 'gsaSupplierContact', 'suppliersData → contact.primary_contact'],
        ],
        header_color='D96F3C'
    )
    
    p = doc.add_paragraph()
    run = p.add_run('Fallback: ')
    run.bold = True
    run = p.add_run('If supplier not found in suppliersData, shows entity from first PO, PO count + total spend, default 4-star rating.')
    run.font.size = Pt(9)
    
    # GSA PO Table
    add_section_header(doc, '7.6 PO Details Table', TAB_COLORS['GSA']['rgb'], level=2)
    
    add_styled_table(doc,
        ['Column', 'Sort Key', 'Field', 'Format'],
        [
            ['PO No.', 'po_no', 'po.poNumber', 'Link'],
            ['Type', 'type', 'po.poType', 'Badge: "CO" (red) / "Base" (green) + group badge'],
            ['Order ID', 'order_id', 'po.orderId', 'Integer sort'],
            ['Project', 'project', 'po.project', 'Truncated 40 chars'],
            ['PO Date', 'po_date', 'po.poDate', 'Date sort (default desc)'],
            ['Supplier', 'supplier', 'po.supplier', 'Text'],
            ['Material', 'material', 'po.material', 'Text'],
            ['PO Value (US$)', 'po_value', 'convertToUSD(po.valueUSD, po.currency)', 'formatCurrency()'],
        ],
        header_color='D96F3C'
    )
    
    p = doc.add_paragraph('Pagination: 25/50/100 rows per page. Default sort: PO Date descending. '
                          'Searchable by PO number, name, project, supplier, material, entity, orderId.')
    p.runs[0].font.size = Pt(9)
    
    # GSA Cross-filter
    add_section_header(doc, '7.7 Cross-Filter Behavior', TAB_COLORS['GSA']['rgb'], level=2)
    
    add_styled_table(doc,
        ['Click Target', 'What Happens'],
        [
            ['Entity bar', 'Filters all components to that entity'],
            ['Project bar', 'Filters all components to that project'],
            ['Top/Bottom supplier bar', 'Updates supplier card + filters to that supplier'],
            ['Trend bar', 'Console log only (no cross-filter)'],
        ],
        header_color='D96F3C'
    )
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════════════
    # 8. M&D TAB
    # ════════════════════════════════════════════════════════════
    
    add_section_header(doc, '8. Tab 3 — Materials & Disciplines (M&D)', TAB_COLORS['MD']['rgb'], level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('Theme: Dark Blue (#0F3D5E)  |  Focus: Material categorization, discipline spend, and conversion analysis')
    run.font.color.rgb = THEME['light']
    run.italic = True
    
    # MD Data Sources
    add_section_header(doc, '8.1 Data Sources', TAB_COLORS['MD']['rgb'], level=2)
    
    add_styled_table(doc,
        ['File', 'Variable', 'Content'],
        [
            ['md_data.json', 'mdData', 'Combined quotations (3,921) + POs (3,746), discipline breakdown, entity breakdown, trend, filters'],
            ['suppliers.json', 'suppliersData', 'Supplier details for profile card and overview table'],
        ],
        header_color='0F3D5E'
    )
    
    # MD Filters
    add_section_header(doc, '8.2 Filters (8 Controls + Search)', TAB_COLORS['MD']['rgb'], level=2)
    
    add_styled_table(doc,
        ['#', 'Filter', 'HTML ID', 'Options Source', 'Matches'],
        [
            ['1', 'Material Code', 'filterMdDiscipline', 'mdData.filters.materialCodes', 'po.materialCode / q.materialCode'],
            ['2', 'Material', 'filterMdMaterial', 'mdData.filters.materials', 'po.material / q.material'],
            ['3', 'Entity', 'filterMdEntity', 'mdData.filters.entities', 'po.entity / q.entity'],
            ['4', 'Project', 'filterMdProject', 'mdData.filters.projects (max 200)', 'po.project / q.project'],
            ['5', 'Supplier', 'filterMdSupplier', 'mdData.filters.suppliers', 'po.supplier / q.supplier'],
            ['6', 'Year', 'filterMdYear', 'Derived from PO years', 'po.year'],
            ['7', 'Date From', 'filterMdFrom', 'Date input', 'poDate / q.date >= fromDate'],
            ['8', 'Date To', 'filterMdTo', 'Date input', 'poDate / q.date <= toDate'],
            ['9', 'Search', 'mdSearchInput', 'Free text (debounced 300ms)', 'Multiple fields'],
        ],
        header_color='0F3D5E'
    )
    
    p = doc.add_paragraph()
    run = p.add_run('Important: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
    run = p.add_run('M&D filters apply to both POs and quotations simultaneously.')
    run.font.size = Pt(9)
    
    # MD KPIs
    add_section_header(doc, '8.3 KPI Cards (5 Metrics)', TAB_COLORS['MD']['rgb'], level=2)
    
    add_styled_table(doc,
        ['#', 'KPI Title', 'Unfiltered Formula', 'Filtered Formula'],
        [
            ['1', 'Materials', 'Unique raw material names (excl. Blank)', 'DISTINCT(po.material + q.material) count'],
            ['2', 'Material Codes', 'Unique consolidated codes', 'DISTINCT(po.materialCode + q.materialCode) count'],
            ['3', 'Total Material Spend', 'SUM(md_pos[].value) in USD', 'SUM(po.value) from filtered POs'],
            ['4', 'Total Material Code Spend', 'Same as #3 (totalOrdered)', 'Same as #3'],
            ['5', 'Active Projects', 'Unique projects from POs', 'DISTINCT(po.project) count'],
        ],
        header_color='0F3D5E'
    )
    
    p = doc.add_paragraph()
    run = p.add_run('Sub-labels: ')
    run.bold = True
    run = p.add_run('KPIs 3 & 4 show conversion %: (totalOrdered / totalQuoted × 100)%. '
                     'KPI 5 shows supplier count: 2,189 from master list.')
    run.font.size = Pt(9)
    run.font.color.rgb = THEME['light']
    
    # MD Charts
    add_section_header(doc, '8.4 Charts', TAB_COLORS['MD']['rgb'], level=2)
    
    md_charts = [
        ['Total Spend by Material Code', 'Horizontal grouped bar', 'disciplineSpendChart',
         'Dataset 1: "Quoted" (#9CB3C9), Dataset 2: "Ordered" (#2B4257). Top 12 by orderedValue.',
         'Unfiltered: mdData.disciplines[]. Filtered: aggregated from filtered POs + quotations by materialCode.'],
        ['Material Distribution', 'Doughnut (55% cutout)', 'materialDistributionChart',
         'Top 10 slices. Click slice → filters by that material.',
         'Unfiltered: mdData.disciplines[] (top 10). Filtered: aggregated from filtered POs by po.material.'],
    ]
    
    add_styled_table(doc,
        ['Chart Name', 'Type', 'Canvas ID', 'Datasets & Interaction', 'Data Source'],
        md_charts,
        header_color='0F3D5E'
    )
    
    # MD Supplier Profile
    add_section_header(doc, '8.5 Supplier Profile Card', TAB_COLORS['MD']['rgb'], level=2)
    
    add_styled_table(doc,
        ['Field', 'HTML ID', 'Source'],
        [
            ['Name', 'mdSupplierName', 'supplier.name'],
            ['Location', 'mdSupplierLocation', 'normalizeCountry() on address data'],
            ['Stars', 'mdSupplierStars', 'rating.score → ⭐/☆'],
            ['Rating', 'mdSupplierRatingVal', 'rating.toFixed(2) + "/5"'],
            ['Email', 'mdSupplierEmail', 'supplier.contact.email'],
            ['Contact', 'mdSupplierContact', 'supplier.contact.primary_contact'],
        ],
        header_color='0F3D5E'
    )
    
    # MD Supplier Table
    add_section_header(doc, '8.6 Supplier Overview Table', TAB_COLORS['MD']['rgb'], level=2)
    
    add_styled_table(doc,
        ['Column', 'Sort', 'Field'],
        [
            ['Supplier Name ↕', 'name', 'supplier.name (clickable → profile)'],
            ['Location ↕', 'country', 'normalizeCountry(address.country_standardized)'],
            ['Rating ↕', 'rating', '⭐ rating.toFixed(1) from suppliers.json'],
            ['Email', '—', 'contact.email'],
            ['Contact', '—', 'contact.primary_contact'],
        ],
        header_color='0F3D5E'
    )
    
    p = doc.add_paragraph('Pagination: 10/25/50 per page with search and page navigation buttons. '
                          'Unfiltered: all suppliers from suppliersData. Filtered: unique suppliers from filtered POs.')
    p.runs[0].font.size = Pt(9)
    
    # MD PO Table
    add_section_header(doc, '8.7 PO/Material Details Table', TAB_COLORS['MD']['rgb'], level=2)
    
    add_styled_table(doc,
        ['Column', 'Field', 'Format'],
        [
            ['PO Number', 'po.poNumber', 'Text'],
            ['PO Date', 'po.poDate', 'Text'],
            ['Material', 'po.material', 'Text'],
            ['Material Code', 'po.materialCode', 'Text'],
            ['PO Value (USD)', 'convertToUSD(po.value, po.currency)', 'formatCurrencyShort()'],
            ['Currency', 'po.currency', 'Text'],
            ['Project', 'po.project', 'Truncated 40 chars'],
        ],
        header_color='0F3D5E'
    )
    
    p = doc.add_paragraph('Pagination: 20 rows/page with prev/next navigation.')
    p.runs[0].font.size = Pt(9)
    
    # MD Approved Materials
    add_section_header(doc, '8.8 Approved Materials', TAB_COLORS['MD']['rgb'], level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('Status: 🚧 Coming Soon ')
    run.bold = True
    run = p.add_run('— Placeholder in current UI. JavaScript functions exist but the HTML tbody '
                     'is hidden by the "Coming Soon" overlay. Will display: Material, Spec No, Supplier, Discipline.')
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════════════
    # 9. DATA FILE REFERENCE
    # ════════════════════════════════════════════════════════════
    
    add_section_header(doc, '9. Data File Reference', THEME['primary'], level=1)
    
    add_section_header(doc, 'Pipeline-Generated Files (7)', THEME['primary'], level=2)
    
    add_styled_table(doc,
        ['File', 'Records', 'Key Fields'],
        [
            ['sm_data.json', '3,921 quotations', '.summary, .workbench[], .statusSummary[], .entities[], .materialsByDiscipline[], .suppliers[] (employees), .filters, .funnel'],
            ['gsa_data.json', '3,746 POs', '.summary, .workbench[], .supplierRankings, .entityBreakdown[], .materialBreakdown[], .monthlyTrend[], .poTypeBreakdown, .changeOrderDetails[], .filters'],
            ['md_data.json', '3,921 Q + 3,746 PO', '.summary, .quotations[], .pos[], .disciplines[], .entityBreakdown[], .trend[], .filters'],
            ['employees.json', '~50 employees', 'Employee quotation/order counts, win rates, spend values'],
            ['conversion_times.json', '180 links', '.monthlyAverage[], .totalLinked, .averageDays'],
            ['change_orders.json', '192 groups', '.totalGroups, .totalCOPOs, .totalCOValue, .groups[]'],
            ['data_metadata.json', '1 record', 'Build date, source files, record counts'],
        ],
        header_color='004578'
    )
    
    add_section_header(doc, 'Pre-existing Reference Files (3)', THEME['primary'], level=2)
    
    add_styled_table(doc,
        ['File', 'Entries', 'Purpose'],
        [
            ['suppliers.json', '2,189 suppliers', 'Master supplier directory with contact, address, rating, material category'],
            ['client_country_map.json', '1,098 entries', 'Client name → country mapping (4-source priority)'],
            ['entity_code_map.json', '~20 entries', 'Entity code → entity name lookup'],
        ],
        header_color='004578'
    )
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════════════
    # 10. MATERIAL CODE CLASSIFICATION
    # ════════════════════════════════════════════════════════════
    
    add_section_header(doc, '10. Material Code Classification', THEME['primary'], level=1)
    
    p = doc.add_paragraph('The pipeline classifies ~35 raw material names into 12 standardized material categories:')
    
    doc.add_paragraph('')
    chart_buf = create_material_codes_chart()
    doc.add_picture(chart_buf, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    
    mat_data = [
        ['Architectural', 'Sandwich Panel, Steel Coil, Doors, Windows, Fit Out Project, Paints, Sanitary Accessories', 'A'],
        ['Chemicals', 'Polyurethane Foam, Chemicals', 'C'],
        ['Electrical', 'Electrical', 'E'],
        ['Fire', 'Firestop/DC 315, Firestop, Fire Alarm, Fire Fighting, Fire Suppression, Fire Protection', 'F'],
        ['Logistics', 'Transportation, Discount, MHE, Logistics', 'L'],
        ['Mechanical', 'Machine/Equipments, Mechanical Items', 'M'],
        ['Office Assets', 'Computer Peripherals', 'O'],
        ['Protection', 'PPE', 'P'],
        ['Rental', 'Rental', 'R'],
        ['Services', 'Design, Construction, LSA - Life Support Area, Subcontract, Services', 'S'],
        ['Tools', 'Tools', 'T'],
        ['Various', 'Containers, Building Materials, Graco Spares, Misc., General', 'V'],
    ]
    
    add_styled_table(doc,
        ['Material Code', 'Raw Materials', 'PO Prefix'],
        mat_data,
        header_color='004578'
    )
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run('PO-based detection: ')
    run.bold = True
    run = p.add_run('When material info is not available from quotation linkage, the pipeline infers material '
                     'from the PO entity code prefix letter (e.g., "M" → Mechanical).')
    run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════════════
    # 11. COUNTRY NORMALIZATION
    # ════════════════════════════════════════════════════════════
    
    add_section_header(doc, '11. Country Normalization', THEME['primary'], level=1)
    
    add_section_header(doc, 'Multi-Source Country Resolution', THEME['primary'], level=2)
    
    add_styled_table(doc,
        ['Priority', 'Source', 'Description'],
        [
            ['1', 'Address', 'Physical address country field'],
            ['2', 'Phone Validation', 'Country from validated phone number'],
            ['3', 'Phone Prefix', 'Country inferred from phone number prefix'],
            ['4', 'Email TLD', 'Country from email domain extension'],
        ],
        header_color='004578'
    )
    
    add_section_header(doc, 'Entity-Based Fallback', THEME['primary'], level=2)
    
    p = doc.add_paragraph('For clients not in the country map, 29 entity→country mappings provide defaults:')
    
    add_styled_table(doc,
        ['Entity', 'Default Country'],
        [
            ['MVL INDUSTRIAL SOLUTIONS', 'United Arab Emirates'],
            ['MVL VENTURES LLC', 'United Arab Emirates'],
            ['MVL ENERGY', 'United Arab Emirates'],
            ['MVL SOLUTIONS', 'United Arab Emirates'],
            ['CENTRICO', 'United Arab Emirates'],
            ['MVL INDUSTRIAL EST', 'Saudi Arabia'],
            ['MVL ARABIA', 'Saudi Arabia'],
            ['MVL PROJECTS', 'United Arab Emirates'],
            ['MVL TRADING', 'United Arab Emirates'],
            ['MVL FACILITIES', 'United Arab Emirates'],
        ],
        header_color='004578'
    )
    
    add_section_header(doc, 'normalizeCountry() Function', THEME['primary'], level=2)
    
    p = doc.add_paragraph('The frontend normalizeCountry() function standardizes ~150 country name variants:')
    
    add_styled_table(doc,
        ['Input Variants', 'Normalized Output'],
        [
            ['"United Arab Emirates" / "UAE" / "U.A.E." / "Uae"', 'United Arab Emirates'],
            ['"Kingdom of Saudi Arabia" / "KSA" / "Saudi"', 'Saudi Arabia'],
            ['"USA" / "United States of America" / "US"', 'United States'],
            ['"UK" / "United Kingdom of Great Britain..."', 'United Kingdom'],
            ['"Peoples Republic of China" / "PRC"', 'China'],
        ],
        header_color='004578'
    )
    
    # ════════════════════════════════════════════════════════════
    # FORMATTING REFERENCE (bonus section)
    # ════════════════════════════════════════════════════════════
    
    doc.add_page_break()
    
    add_section_header(doc, 'Appendix — Design Tokens & Formatting', THEME['primary'], level=1)
    
    add_section_header(doc, 'Currency Display Functions', THEME['primary'], level=2)
    
    add_styled_table(doc,
        ['Context', 'Function', 'Example'],
        [
            ['KPI cards', 'formatCurrencyShort()', '$478.6M, $12.1K, $1.2B'],
            ['Table cells', 'formatCurrency()', '$1.2M, $478.6K, $12'],
            ['Number only', 'formatNumber()', '3,746'],
        ],
        header_color='004578'
    )
    
    add_section_header(doc, 'Status Badge Colors', THEME['primary'], level=2)
    
    add_styled_table(doc,
        ['Status', 'Color', 'Hex'],
        [
            ['Order', 'Green', '#2ECC71'],
            ['Quotation', 'Blue', '#3498DB'],
            ['Waiting', 'Orange', '#F39C12'],
            ['Cancelled', 'Red', '#E74C3C'],
            ['Closed', 'Gray', '#95A5A6'],
        ],
        header_color='004578'
    )
    
    add_section_header(doc, 'Change Order Badges', THEME['primary'], level=2)
    
    add_styled_table(doc,
        ['Badge', 'Color', 'Meaning'],
        [
            ['CO', 'Red #E74C3C', 'Change Order'],
            ['Base', 'Green #2ECC71', 'Base PO'],
            ['N of M', 'Gold #F39C12', 'Group indicator (e.g., "2 of 3")'],
        ],
        header_color='004578'
    )
    
    # ── Footer note ──
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 50)
    run.font.color.rgb = THEME['light']
    run.font.size = Pt(8)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Document generated from codebase analysis of v8/shared/scripts.js (~5,928 lines),\n'
                     'v8/data/build_v8_data.py (~1,508 lines), and v8/index.html (~1,063 lines).\n'
                     'MVL Supply Chain Intelligence Hub — V8 — February 2026')
    run.font.size = Pt(8)
    run.font.color.rgb = THEME['light']
    
    # ── Save ──
    doc.save(OUTPUT_PATH)
    print(f"\n✅ Document saved: {OUTPUT_PATH}")
    print(f"   Pages: ~25+ pages with charts, tables, and diagrams")
    return doc

# ═══════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════

if __name__ == '__main__':
    print("=" * 60)
    print("  MVL Dashboard — DOCX Documentation Generator")
    print("=" * 60)
    build_document()
